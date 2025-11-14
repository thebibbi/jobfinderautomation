from typing import Dict, Any
from loguru import logger
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from pathlib import Path
import json

from .claude_service import get_claude_service
from ..prompts.cover_letter import COVER_LETTER_PROMPT


class DocumentGenerator:
    """Generates tailored resumes and cover letters"""

    def __init__(self):
        self.claude = get_claude_service()

    async def generate_resume(
        self,
        job_data: Dict[str, Any],
        analysis_results: Dict[str, Any]
    ) -> bytes:
        """
        Generate tailored resume as DOCX

        Uses building blocks + Claude to create optimized resume
        """
        logger.info(f"📄 Generating resume for: {job_data['company']} - {job_data['job_title']}")

        try:
            # Extract key information from analysis
            recommended_skills = analysis_results.get("resume_tailoring", {}).get("skills_to_highlight", [])
            must_include_achievements = analysis_results.get("resume_tailoring", {}).get("must_include_achievements", [])
            suggested_summary = analysis_results.get("resume_tailoring", {}).get("suggested_summary", "")

            # Create Word document
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Header - Name and Contact
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = header.add_run("[Candidate Name]\n")
            name_run.font.size = Pt(16)
            name_run.font.bold = True

            contact_run = header.add_run(
                "Email: email@example.com | Phone: (555) 123-4567 | "
                "LinkedIn: linkedin.com/in/profile\n"
            )
            contact_run.font.size = Pt(10)

            # Professional Summary
            doc.add_heading('Professional Summary', level=1)
            summary_text = suggested_summary if suggested_summary else "Data-driven professional transitioning from education to corporate analytics. Strong background in mathematics, data analysis, and project management."
            doc.add_paragraph(summary_text)

            # Core Competencies
            doc.add_heading('Core Competencies', level=1)
            skills_para = doc.add_paragraph()
            skills_text = " • ".join(recommended_skills[:10] if recommended_skills else ["Data Analysis", "Project Management", "Stakeholder Communication"])
            skills_para.add_run(skills_text)

            # Professional Experience
            doc.add_heading('Professional Experience', level=1)

            # Add sample experience
            job_para = doc.add_paragraph()
            job_run = job_para.add_run("Mathematics Educator & Data Analyst | School District\n")
            job_run.font.bold = True
            date_run = job_para.add_run("2014 - 2024")
            date_run.font.italic = True

            # Add achievements
            for achievement in must_include_achievements[:5]:
                doc.add_paragraph(achievement, style='List Bullet')

            # Education
            doc.add_heading('Education', level=1)
            doc.add_paragraph(
                "Master of Science in Mathematics\n[University Name] | [Year]",
                style='List Bullet'
            )

            # Save to bytes
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            logger.info("✅ Resume generated successfully")
            return doc_bytes.read()

        except Exception as e:
            logger.error(f"❌ Error generating resume: {e}")
            raise

    async def generate_cover_letter(
        self,
        job_data: Dict[str, Any],
        analysis_results: Dict[str, Any],
        style: str = "conversational"
    ) -> str:
        """
        Generate tailored cover letter

        Args:
            job_data: Job information
            analysis_results: Analysis from Claude
            style: "conversational" or "formal"
        """
        logger.info(f"📝 Generating {style} cover letter for: {job_data['company']}")

        try:
            # Load voice profile
            voice_profile_path = Path(__file__).parent.parent.parent.parent / "skills" / "voice_profile.md"
            with open(voice_profile_path, 'r') as f:
                voice_profile = f.read()

            # Get cover letter guidance from analysis
            guidance = analysis_results.get("cover_letter_guidance", {})

            # Build comprehensive prompt
            full_prompt = f"""
{voice_profile}

<job_details>
<company>{job_data['company']}</company>
<job_title>{job_data['job_title']}</job_title>
<job_description>
{job_data['job_description']}
</job_description>
</job_details>

<analysis_guidance>
{json.dumps(guidance, indent=2)}
</analysis_guidance>

<style>{style}</style>

{COVER_LETTER_PROMPT}
"""

            response = self.claude.client.messages.create(
                model=self.claude.model,
                max_tokens=self.claude.max_tokens,
                messages=[{
                    "role": "user",
                    "content": full_prompt
                }]
            )

            cover_letter = response.content[0].text

            logger.info(f"✅ {style.title()} cover letter generated")
            return cover_letter

        except Exception as e:
            logger.error(f"❌ Error generating cover letter: {e}")
            raise


# Singleton
_document_generator = None

def get_document_generator() -> DocumentGenerator:
    global _document_generator
    if _document_generator is None:
        _document_generator = DocumentGenerator()
    return _document_generator
