"""
Document Conversion Service

Converts various document formats (DOCX, PDF, MD) to Markdown format
for efficient LLM processing with minimal token usage.
"""

from typing import Optional
from pathlib import Path
import io
from loguru import logger
from docx import Document
import PyPDF2


class DocumentConverter:
    """Converts documents to Markdown format for LLM efficiency"""

    @staticmethod
    def docx_to_markdown(file_content: bytes) -> str:
        """
        Convert DOCX file to Markdown

        Args:
            file_content: DOCX file as bytes

        Returns:
            Markdown formatted string
        """
        try:
            logger.info("📄 Converting DOCX to Markdown")

            # Load DOCX from bytes
            doc = Document(io.BytesIO(file_content))

            markdown_lines = []

            # Extract paragraphs and convert to markdown
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()

                if not text:
                    # Preserve paragraph spacing
                    markdown_lines.append("")
                    continue

                # Check if it's a heading (bold and/or larger font)
                if paragraph.style.name.startswith('Heading'):
                    # Determine heading level (Heading 1 -> #, Heading 2 -> ##, etc.)
                    level = paragraph.style.name.replace('Heading ', '')
                    try:
                        level_num = int(level)
                        markdown_lines.append(f"{'#' * level_num} {text}")
                    except ValueError:
                        markdown_lines.append(f"## {text}")
                else:
                    markdown_lines.append(text)

            # Extract tables
            for table in doc.tables:
                markdown_lines.append("")  # Spacing before table

                # Get column count
                if table.rows:
                    col_count = len(table.rows[0].cells)

                    # Header row
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    markdown_lines.append("| " + " | ".join(header_cells) + " |")
                    markdown_lines.append("| " + " | ".join(["---"] * col_count) + " |")

                    # Data rows
                    for row in table.rows[1:]:
                        cells = [cell.text.strip() for cell in row.cells]
                        markdown_lines.append("| " + " | ".join(cells) + " |")

                markdown_lines.append("")  # Spacing after table

            markdown_text = "\n".join(markdown_lines)

            # Clean up excessive blank lines
            while "\n\n\n" in markdown_text:
                markdown_text = markdown_text.replace("\n\n\n", "\n\n")

            logger.info(f"✅ DOCX converted to Markdown ({len(markdown_text)} chars)")
            return markdown_text.strip()

        except Exception as e:
            logger.error(f"❌ Error converting DOCX to Markdown: {e}")
            raise ValueError(f"Failed to convert DOCX: {str(e)}")

    @staticmethod
    def pdf_to_markdown(file_content: bytes) -> str:
        """
        Convert PDF file to Markdown

        Args:
            file_content: PDF file as bytes

        Returns:
            Markdown formatted string
        """
        try:
            logger.info("📄 Converting PDF to Markdown")

            # Read PDF from bytes
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

            markdown_lines = []

            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()

                if text.strip():
                    # Add page separator for multi-page PDFs
                    if page_num > 1:
                        markdown_lines.append(f"\n---\n**Page {page_num}**\n")

                    # Clean up the extracted text
                    lines = text.split('\n')
                    processed_lines = []

                    for line in lines:
                        line = line.strip()
                        if line:
                            # Try to detect headings (all caps lines)
                            if line.isupper() and len(line.split()) <= 8:
                                processed_lines.append(f"\n## {line.title()}\n")
                            else:
                                processed_lines.append(line)

                    markdown_lines.extend(processed_lines)

            markdown_text = "\n".join(markdown_lines)

            # Clean up excessive blank lines
            while "\n\n\n" in markdown_text:
                markdown_text = markdown_text.replace("\n\n\n", "\n\n")

            logger.info(f"✅ PDF converted to Markdown ({len(markdown_text)} chars)")
            return markdown_text.strip()

        except Exception as e:
            logger.error(f"❌ Error converting PDF to Markdown: {e}")
            raise ValueError(f"Failed to convert PDF: {str(e)}")

    @staticmethod
    def validate_markdown(file_content: bytes) -> str:
        """
        Validate and clean Markdown file

        Args:
            file_content: Markdown file as bytes

        Returns:
            Cleaned Markdown string
        """
        try:
            logger.info("📄 Validating Markdown file")

            # Decode bytes to string
            markdown_text = file_content.decode('utf-8')

            # Basic cleanup
            markdown_text = markdown_text.strip()

            # Clean up excessive blank lines
            while "\n\n\n" in markdown_text:
                markdown_text = markdown_text.replace("\n\n\n", "\n\n")

            logger.info(f"✅ Markdown validated ({len(markdown_text)} chars)")
            return markdown_text

        except UnicodeDecodeError as e:
            logger.error(f"❌ Invalid UTF-8 encoding in Markdown file: {e}")
            raise ValueError("Markdown file must be UTF-8 encoded")
        except Exception as e:
            logger.error(f"❌ Error validating Markdown: {e}")
            raise ValueError(f"Failed to validate Markdown: {str(e)}")

    def convert_to_markdown(
        self,
        file_content: bytes,
        filename: str
    ) -> str:
        """
        Convert any supported file format to Markdown

        Args:
            file_content: File content as bytes
            filename: Original filename (used to determine type)

        Returns:
            Markdown formatted string

        Raises:
            ValueError: If file format is not supported
        """
        # Determine file type from extension
        file_ext = Path(filename).suffix.lower()

        if file_ext == '.docx':
            return self.docx_to_markdown(file_content)
        elif file_ext == '.pdf':
            return self.pdf_to_markdown(file_content)
        elif file_ext in ['.md', '.markdown']:
            return self.validate_markdown(file_content)
        else:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                "Supported formats: .docx, .pdf, .md"
            )


# Singleton instance
_converter = None

def get_document_converter() -> DocumentConverter:
    """Get singleton DocumentConverter instance"""
    global _converter
    if _converter is None:
        _converter = DocumentConverter()
    return _converter
